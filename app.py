from flask import Flask, render_template, request, redirect, url_for, flash,send_file
import mysql.connector
import secrets
from mysql.connector import Error
from docx import Document
from io import BytesIO
from flask import session
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import random
from flask_mail import Mail, Message
app = Flask(__name__)
app.secret_key = secrets.token_hex(16)

# Database configuration
DB_CONFIG = {
    "host": "localhost",
    "user": "root",
    "password": "Pass@123",
    "database": "question_bank",
}

# Establish database connection
def create_connection():
    try:
        return mysql.connector.connect(**DB_CONFIG)
    except Error as e:
        print(f"Error connecting to MySQL: {e}")
        return None

# Create tables
def initialize_database():
    connection = create_connection()
    if connection:
        try:
            cursor = connection.cursor()
            cursor.execute('''

            CREATE TABLE users (
                id INT AUTO_INCREMENT PRIMARY KEY,
                email VARCHAR(255) UNIQUE NOT NULL,
                username VARCHAR(255) UNIQUE NOT NULL,
                password VARCHAR(255) NOT NULL,
                role VARCHAR(50) NOT NULL,
                otp VARCHAR(50) 
            );

        ''')

            # Create Subjects table
            create_subjects_table = """
            CREATE TABLE IF NOT EXISTS subjects (
                subject_id INT AUTO_INCREMENT PRIMARY KEY,
                subject_name VARCHAR(100) NOT NULL,
                branch VARCHAR(100) NOT NULL,
                UNIQUE (subject_name, branch)
            );
            """
            cursor.execute(create_subjects_table)

            # Create Questions table
            create_questions_table = """
            CREATE TABLE IF NOT EXISTS questions (
                question_id INT AUTO_INCREMENT PRIMARY KEY,
                question_text TEXT NOT NULL,
                module INT NOT NULL,
                difficulty ENUM('easy', 'medium', 'hard') NOT NULL,
                subject_id INT,
                FOREIGN KEY (subject_id) REFERENCES subjects(subject_id) ON DELETE CASCADE
            );
            """
            cursor.execute(create_questions_table)

            print("Tables created successfully!")
        except Error as e:
            print(f"Error creating tables: {e}")
        finally:
            cursor.close()
            connection.close()
            
 # Routes
@app.route('/')
def index():
    return render_template('index.html')


#signup
@app.route('/signup', methods=['GET', 'POST'])
def signup():
    if request.method == 'POST':
        email = request.form.get('email')
        username = request.form.get('username')
        password = request.form.get('password')
        confirm_password = request.form.get('confirm-password')
        role = request.form.get('role')  # Retrieve role from the form

        # Check for missing fields
        if not email or not username or not password or not confirm_password or not role:
            flash('All fields are required!', 'error')
            return redirect(url_for('signup'))

        # Check if passwords match
        if password != confirm_password:
            flash('Passwords do not match!', 'error')
            return redirect(url_for('signup'))

        try:
            conn = mysql.connector.connect(**DB_CONFIG)
            cursor = conn.cursor()

            # Include the role field in the INSERT statement
            cursor.execute('''
                INSERT INTO users (email, username, password, role) 
                VALUES (%s, %s, %s, %s)
            ''', (email, username, password, role))
            
            conn.commit()
            cursor.close()
            conn.close()

            flash('Sign-up successful! Please log in.', 'success')
            return redirect(url_for('login'))
        except mysql.connector.Error as e:
            flash('Email or Username already exists!', 'error')
            return redirect(url_for('signup'))

    return render_template('signup.html')


#login
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form.get('username')  
        password = request.form.get('password')

        if not username or not password:
            flash('All fields are required!', 'error')
            return render_template('login.html')  

        try:
            conn = mysql.connector.connect(**DB_CONFIG)
            cursor = conn.cursor()
            cursor.execute('SELECT * FROM users WHERE username = %s AND password = %s', (username, password))
            user = cursor.fetchone()
            cursor.close()
            conn.close()

            if user:
                session['username'] = user[2]  # Assuming username is the 3rd column
                session['role'] = user[4]      # Assuming role is the 5th column
                flash('Login successful!', 'success')

                if session['role'] == 'Teacher':
                    return redirect(url_for('add_question'))
                else:
                    return redirect(url_for('generate_question_paper'))
            else:
                flash('Invalid username or password!', 'error')
                return render_template('login.html')  

        except mysql.connector.Error as e:
            print("Database Error:", e)
            flash('An error occurred while connecting to the database.', 'error')
            return render_template('login.html')  

    return render_template('login.html')


app.config['MAIL_SERVER'] = 'smtp.gmail.com'  # Use your SMTP server details
app.config['MAIL_PORT'] = 465
app.config['MAIL_USE_SSL'] = True
app.config['MAIL_DEFAULT_SENDER'] = 'vedikapatil2713@gmail.com'
app.config['MAIL_USERNAME'] = 'vedikapatil2713@gmail.com'  # Your email address
app.config['MAIL_PASSWORD'] = 'iemz htjj jrxz thfc'  # Your email password
mail = Mail(app)

app.config['MAIL_SERVER'] = 'smtp.gmail.com'  # Use your SMTP server details
app.config['MAIL_PORT'] = 465
app.config['MAIL_USE_SSL'] = True
app.config['MAIL_DEFAULT_SENDER'] = 'srushtibhosale436@gmail.com'
app.config['MAIL_USERNAME'] = 'srushtibhosale436@gmail.com'  # Your email address
app.config['MAIL_PASSWORD'] = 'iemz htjj jrxz thfc'  # Your email password
mail = Mail(app)

# Route to display the "Forgot Password" form
@app.route('/forgot_password', methods=['GET', 'POST'])
def forgot_password():
    if request.method == 'POST':
        email = request.form.get('email')
        
        # Check if the email exists in the database
        connection = create_connection()
        if connection:
            try:
                cursor = connection.cursor()
                cursor.execute("SELECT * FROM users WHERE email = %s", (email,))
                user = cursor.fetchone()

                if user:
                    # Generate OTP using pyotp
                    otp = f"{random.randint(1000, 9999)}" # Generate a random OTP

                    # Store the OTP in the database temporarily (you can store it for a limited time in your real application)
                    cursor.execute("UPDATE users SET otp = %s WHERE email = %s", (otp, email))
                    connection.commit()

                    # Send the OTP to the user's email
                    msg = Message("Your OTP for Password Reset", recipients=[email])
                    msg.body = f"Your OTP is: {otp}. Use this OTP to reset your password."
                    mail.send(msg)

                    flash('OTP sent to your email address.', 'info')
                    return redirect(url_for('verify_otp', email=email))  # Redirect to OTP verification page
                else:
                    flash('Email not found!', 'error')

            except Error as e:
                flash(f"Error: {e}", 'error')
            finally:
                cursor.close()
                connection.close()
        else:
            flash('Failed to connect to the database.', 'error')
        
    return render_template('forgot_password.html')


# Route to verify OTP and reset the password
@app.route('/verify_otp', methods=['GET', 'POST'])
def verify_otp():
    email = request.args.get('email')  # Get email from the URL
    if request.method == 'POST':
        otp = request.form.get('otp')
        new_password = request.form.get('new_password')
        confirm_password = request.form.get('confirm_password')

        if not otp or not new_password or not confirm_password:
            flash('All fields are required!', 'error')
            return redirect(url_for('verify_otp', email=email))

        if new_password != confirm_password:
            flash('Passwords do not match!', 'error')
            return redirect(url_for('verify_otp', email=email))

        # Validate OTP and reset the password
        connection = create_connection()
        if connection:
            try:
                cursor = connection.cursor()
                cursor.execute("SELECT otp FROM users WHERE email = %s", (email,))
                stored_otp = cursor.fetchone()

                if stored_otp and stored_otp[0] == otp:
                    # Update the user's password in the database
                    cursor.execute("UPDATE users SET password = %s WHERE email = %s", (new_password, email))
                    connection.commit()
                    flash('Password reset successful!', 'success')
                    return redirect(url_for('login'))
                else:
                    flash('Invalid OTP!', 'error')

            except Error as e:
                flash(f"Error: {e}", 'error')
            finally:
                cursor.close()
                connection.close()
        else:
            flash('Failed to connect to the database.', 'error')

    return render_template('verify_otp.html', email=email)


            
#add_question
# Route to render the HTML form and handle form submissions
@app.route('/add_question', methods=["GET", "POST"])
def add_question():
    if 'username' not in session:
        flash("Please log in to access this page.", "warning")
        return redirect(url_for('login'))
    
    if request.method == "POST":
        # Retrieve form data
        question_text = request.form["question_text"]
        difficulty = request.form["difficulty"]
        module = int(request.form["module"])
        subject = request.form["subject"]
        branch = request.form["branch"]

        connection = create_connection()
        if connection:
            try:
                cursor = connection.cursor()

                # Check if the subject exists in the subjects table
                cursor.execute(
                    "SELECT subject_id FROM subjects WHERE subject_name = %s AND branch = %s",
                    (subject, branch),
                )
                result = cursor.fetchone()

                # If subject doesn't exist, insert it
                if not result:
                    cursor.execute(
                        "INSERT INTO subjects (subject_name, branch) VALUES (%s, %s)",
                        (subject, branch),
                    )
                    connection.commit()
                    subject_id = cursor.lastrowid
                else:
                    subject_id = result[0]

                # Insert question into the questions table
                cursor.execute(
                    """
                    INSERT INTO questions (question_text, module, difficulty, subject_id)
                    VALUES (%s, %s, %s, %s)
                    """,
                    (question_text, module, difficulty, subject_id),
                )
                connection.commit()

                flash("Question added successfully!")
                return redirect(url_for('add_question'))
            except Error as e:
                return f"Error: {e}"
            finally:
                cursor.close()
                connection.close()
        else:
            return "Failed to connect to the database."

    return render_template("add_question.html")

@app.route("/show_questions")
def show_questions():
    connection = create_connection()
    questions_by_subject = {}
    if connection:
        try:
            cursor = connection.cursor()

            # Get all questions with their subject name
            cursor.execute("""
            SELECT q.question_text, q.module, q.difficulty, s.subject_name
            FROM questions q
            JOIN subjects s ON q.subject_id = s.subject_id
            """)
            questions = cursor.fetchall()

            # Group questions by subject
            for question in questions:
                subject_name = question[3]
                if subject_name not in questions_by_subject:
                    questions_by_subject[subject_name] = []
                questions_by_subject[subject_name].append(question)

        except Error as e:
            print(f"Error fetching questions: {e}")
        finally:
            cursor.close()
            connection.close()

    return render_template("show_questions.html", questions_by_subject=questions_by_subject)  # Render show questions page

@app.route('/generate_question_paper', methods=['GET', 'POST'])
def generate_question_paper():
    if request.method == 'POST':
        # Extract form data
        branch = request.form.get('branch')
        semester = request.form.get('semester')
        subject = request.form.get('subject')
        subject_code = request.form.get('subject_code')  # New field for subject code
        difficulty = request.form.get('difficulty')
        paper_name = request.form.get('paper_name')
        paper_code = request.form.get('paper_code')
        marks_per_module = request.form.get('marks_per_module')
        total_marks = request.form.get('total_marks')
        exam_time = request.form.get('exam_time')
        exam_date = request.form.get('exam_date')
        examination = request.form.get('examination')
        choice_per_module = request.form.get('choice_per_module')  # New field for choice per module

        # Validate numeric inputs
        try:
            modules = int(request.form.get('modules'))
            questions_per_module = int(request.form.get('questions_per_module'))
        except (TypeError, ValueError):
            flash('Modules and Questions Per Module must be valid numbers!', 'error')
            return redirect(url_for('generate_question_paper'))

        # Validate required fields
        if not all([branch, semester, subject, subject_code, difficulty, paper_name, paper_code, marks_per_module, total_marks, exam_time, exam_date, examination, choice_per_module]):
            flash('All fields are required!', 'error')
            return redirect(url_for('generate_question_paper'))

        connection = create_connection()
        if connection:
            try:
                cursor = connection.cursor()

                # Create a Word document
                doc = Document()

                # Add "Seat No." and "QP Code" section
                p = doc.add_paragraph()
                p.add_run("Seat No.: ").bold = True
                p.add_run(" " * 100)  # Spacer
                p.add_run(f"QP Code: {paper_code}").bold = True

                # Add examination name
                exam_heading = doc.add_paragraph(examination)
                exam_heading.alignment = 1  # Center alignment
                exam_heading.runs[0].font.size = Pt(18)
                exam_heading.runs[0].bold = True

                # Add subject details
                subject_line = doc.add_paragraph()
                subject_line.add_run(f"Subject Name: {paper_name}").bold = True
                subject_line.add_run(" " * 50)  # Spacer
                subject_line.add_run(f"Subject Code: {subject_code}").bold = True

                # Add exam date and total marks
                date_marks_line = doc.add_paragraph()
                date_marks_line.add_run(f"Day & Date: {exam_date}").bold = True
                date_marks_line.add_run(" " * 70)  # Spacer
                date_marks_line.add_run(f"Total Marks: {total_marks}").bold = True

                # Add time
                time_line = doc.add_paragraph()
                time_line.add_run(f"Time: {exam_time}").bold = True

                # Add instructions
                instructions = doc.add_paragraph("Instructions:")
                instructions.runs[0].bold = True
                instructions.add_run("\n1) All questions are compulsory.")

                # Add a horizontal line (broad line separator)
                doc.add_paragraph("=" * 50)

                # Add the questions for each module
                for module in range(1, modules + 1):
                    # Fetch questions for the module
                    query = """
                    SELECT question_text
                    FROM questions q
                    JOIN subjects s ON q.subject_id = s.subject_id
                    WHERE s.branch = %s AND s.subject_name = %s AND q.module = %s 
                    """
                    cursor.execute(query, (branch, subject, module))
                    questions = cursor.fetchall()

                    if questions:
                        # Store questions in an array and shuffle it randomly
                        question_array = [question[0] for question in questions]
                        random.shuffle(question_array)

                        # Add module heading with choice per module
                        module_heading = doc.add_paragraph(f" {module}.Answer The following Question  (Any {choice_per_module})- {marks_per_module} marks")
                        module_heading.style = 'Heading 3'

                        # Add questions from the shuffled array to the doc
                        for i, question_text in enumerate(question_array[:questions_per_module], start=1):
                            doc.add_paragraph(f"{i}. {question_text}")
                    else:
                        doc.add_paragraph(f"No questions available for Module {module}.")


                # Save the document in memory
                file_stream = BytesIO()
                doc.save(file_stream)
                file_stream.seek(0)

                # Send the Word file as a response
                return send_file(file_stream, as_attachment=True, download_name=f"{paper_name}.docx")

            except mysql.connector.Error as e:
                flash(f"Database error: {e}", 'error')
            finally:
                cursor.close()
                connection.close()

    return render_template('generate_question_paper.html')





@app.route('/logout')
def logout():
    # Clear the session to log out the user
    session.clear()
    flash('You have been logged out successfully.', 'info')
    return redirect(url_for('login'))


 
# Initialize database and tables
initialize_database()

if __name__ == "__main__":
    app.run(debug=True)
    
